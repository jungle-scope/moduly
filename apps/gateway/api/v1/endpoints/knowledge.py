import logging
import mimetypes
import os
import tempfile
from typing import List
from uuid import UUID

import pandas as pd
import requests
from docx import Document as DocxDocument
from fastapi import (
    APIRouter,
    BackgroundTasks,
    Depends,
    HTTPException,
    Response,
    status,
)
from fastapi.responses import (
    FileResponse,
    HTMLResponse,
    RedirectResponse,
    StreamingResponse,
)
from sqlalchemy import func
from sqlalchemy.orm import Session

from apps.gateway.api.deps import get_db
from apps.gateway.auth.dependencies import get_current_user
from apps.gateway.services.ingestion.service import (
    IngestionOrchestrator as IngestionService,
)
from apps.shared.db.models.knowledge import Document, KnowledgeBase
from apps.shared.db.models.user import User
from apps.shared.schemas.rag import (
    DocumentPreviewRequest,
    DocumentPreviewResponse,
    DocumentResponse,
    KnowledgeBaseCreate,
    KnowledgeBaseDetailResponse,
    KnowledgeBaseResponse,
    KnowledgeUpdate,
)

logger = logging.getLogger(__name__)
router = APIRouter()


@router.post(
    "", response_model=KnowledgeBaseResponse, status_code=status.HTTP_201_CREATED
)
def create_knowledge_base(
    kb_in: KnowledgeBaseCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    빈 지식 베이스를 생성합니다. (소스 없음)
    """
    # 임베딩 모델 유효성 검사 등은 생략하거나 추후 추가
    kb = KnowledgeBase(
        name=kb_in.name,
        description=kb_in.description,
        embedding_model=kb_in.embedding_model,
        user_id=current_user.id,
    )
    db.add(kb)
    db.commit()
    db.refresh(kb)

    return KnowledgeBaseResponse(
        id=kb.id,
        name=kb.name,
        description=kb.description,
        document_count=0,
        created_at=kb.created_at,
        updated_at=kb.updated_at,
        source_types=[],
        embedding_model=kb.embedding_model,
    )


# TODO: is_active column 추가해서 LLM노드와 지식 베이스 목록에서 모두 사용할 수 있도록 한다
@router.get("", response_model=List[KnowledgeBaseResponse])
def list_knowledge_bases(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    사용자의 자료 목록을 조회합니다.
    각 지식 베이스 그룹에 포함된 문서 개수도 함께 반환합니다.
    """
    # 완료된 문서만 카운트
    # completed_count = func.sum(
    #     case((Document.status == "completed", 1), else_=0)
    # ).label("document_count")

    results = (
        db.query(
            KnowledgeBase,
            func.count(Document.id).label("document_count"),
            func.max(Document.updated_at).label("last_updated_at"),
            func.array_agg(Document.source_type).label("source_types"),
        )
        .outerjoin(Document, KnowledgeBase.id == Document.knowledge_base_id)
        .filter(KnowledgeBase.user_id == current_user.id)
        .group_by(KnowledgeBase.id)
        .order_by(KnowledgeBase.created_at.desc())
        .all()
    )

    response = []
    for kb, doc_count, last_updated_at, source_types in results:
        # source_types가 [None]인 경우 (문서가 없을 때) 빈 리스트로 처리
        clean_source_types = [st for st in source_types if st is not None]

        # KB 업데이트 시간과 문서 최신 업데이트 시간 중 더 최신을 선택
        # 문서가 없으면 KB 업데이트 시간 사용
        final_updated_at = (
            max(kb.updated_at, last_updated_at) if last_updated_at else kb.updated_at
        )

        response.append(
            KnowledgeBaseResponse(
                id=kb.id,
                name=kb.name,
                description=kb.description,
                document_count=doc_count,
                created_at=kb.created_at,
                updated_at=final_updated_at,
                source_types=clean_source_types,
                embedding_model=kb.embedding_model,
            )
        )
    return response


@router.get("/{kb_id}", response_model=KnowledgeBaseDetailResponse)
def get_knowledge_base(
    kb_id: UUID,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    지식 베이스의 상세 정보를 조회합니다.
    포함된 자료 목록과 각 자료의 상태를 함께 반환합니다.
    """
    kb = (
        db.query(KnowledgeBase)
        .filter(KnowledgeBase.id == kb_id, KnowledgeBase.user_id == current_user.id)
        .first()
    )

    if not kb:
        raise HTTPException(status_code=404, detail="Knowledge Base not found")

    # 문서 목록 변환
    doc_responses = []
    for doc in kb.documents:
        # TODO: 청크 개수나 토큰 수는 별도 쿼리로 최적화 필요 (현재는 Lazy Loading)
        doc_responses.append(
            DocumentResponse(
                id=doc.id,
                filename=doc.filename,
                status=doc.status,
                created_at=doc.created_at,
                updated_at=doc.updated_at,
                error_message=doc.error_message,
                chunk_count=len(doc.chunks),  # N+1 발생 가능, 추후 최적화
                token_count=0,  # 우선 0으로 반환
                source_type=doc.source_type,
                meta_info=doc.meta_info,
            )
        )

    return KnowledgeBaseDetailResponse(
        id=kb.id,
        name=kb.name,
        description=kb.description,
        document_count=len(doc_responses),
        created_at=kb.created_at,
        embedding_model=kb.embedding_model,
        documents=doc_responses,
    )


@router.patch("/{kb_id}", status_code=status.HTTP_204_NO_CONTENT)
def update_knowledge_base(
    kb_id: UUID,
    update_data: KnowledgeUpdate,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    지식 베이스의 설정을 수정합니다. (이름, 설명, 즐겨찾기 임베딩 모델)
    """
    kb = (
        db.query(KnowledgeBase)
        .filter(KnowledgeBase.id == kb_id, KnowledgeBase.user_id == current_user.id)
        .first()
    )

    if not kb:
        raise HTTPException(status_code=404, detail="Knowledge Base not found")

    if update_data.name is not None:
        kb.name = update_data.name
    if update_data.description is not None:
        kb.description = update_data.description

    # 임베딩 모델 변경 및 재인덱싱 트리거
    if (
        update_data.embedding_model is not None
        and update_data.embedding_model != kb.embedding_model
    ):
        kb.embedding_model = update_data.embedding_model

        # 재인덱싱 트리거
        orchestrator = IngestionService(db, current_user.id)
        background_tasks.add_task(
            orchestrator.reindex_knowledge_base, kb.id, update_data.embedding_model
        )

    db.commit()
    db.refresh(kb)

    return Response(status_code=status.HTTP_204_NO_CONTENT)


@router.delete("/{kb_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_knowledge_base(
    kb_id: UUID,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    지식 베이스를 삭제합니다.
    연결된 문서 및 임베딩 데이터는 DB Cascade 설정에 따라 함께 삭제됩니다.
    물리적 파일(S3/Local)도 함께 삭제합니다.
    """
    kb = (
        db.query(KnowledgeBase)
        .filter(KnowledgeBase.id == kb_id, KnowledgeBase.user_id == current_user.id)
        .first()
    )

    if not kb:
        raise HTTPException(status_code=404, detail="Knowledge Base not found")

    # 물리적 파일 삭제 (Storage)
    from services.storage import get_storage_service

    storage = get_storage_service()

    for doc in kb.documents:
        if doc.file_path:
            try:
                # S3/Local 파일 삭제
                storage.delete(doc.file_path)
            except Exception as e:
                # 파일 삭제 실패하더라도 DB 삭제는 계속 진행 (로그만 남김)
                logger.warning(
                    f"Failed to delete file {doc.file_path} for doc {doc.id}: {e}"
                )

    # DB 삭제 (Cascade로 청크도 같이 삭제됨)
    db.delete(kb)
    db.commit()

    return Response(status_code=status.HTTP_204_NO_CONTENT)


@router.get("/{kb_id}/documents/{document_id}", response_model=DocumentResponse)
def get_document(
    kb_id: UUID,
    document_id: UUID,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    특정 문서를 조회합니다.
    """
    doc = (
        db.query(Document)
        .join(KnowledgeBase)
        .filter(Document.id == document_id, KnowledgeBase.user_id == current_user.id)
        .first()
    )

    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    if doc.knowledge_base_id != kb_id:
        raise HTTPException(
            status_code=400, detail="Document does not belong to this Knowledge Base"
        )

    return DocumentResponse(
        id=doc.id,
        filename=doc.filename,
        status=doc.status,
        created_at=doc.created_at,
        updated_at=doc.updated_at,
        error_message=doc.error_message,
        chunk_count=len(doc.chunks),
        # token_count=doc.token_count,
        source_type=doc.source_type,
        meta_info=doc.meta_info,
    )


@router.get("/{kb_id}/documents/{document_id}/content")
def get_document_content(
    kb_id: UUID,
    document_id: UUID,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    문서의 원본 파일을 반환합니다. (브라우저 표시용)
    """
    doc = (
        db.query(Document)
        .join(KnowledgeBase)
        .filter(Document.id == document_id, KnowledgeBase.user_id == current_user.id)
        .first()
    )

    if not doc:
        raise HTTPException(status_code=404, detail="Document not found in DB")

    if doc.knowledge_base_id != kb_id:
        raise HTTPException(
            status_code=400, detail="Document does not belong to this Knowledge Base"
        )

    # API 소스는 파일이 없으므로 미리보기 불가 처리
    if doc.source_type == "API" or not doc.file_path:
        raise HTTPException(
            status_code=400,
            detail="API로 받은 응답은 원문 보기를 제공하지 않습니다.",
        )

    # file path가 S3 URL인지 확인합니다.
    is_s3_file = str(doc.file_path).startswith("http") or str(doc.file_path).startswith(
        "s3://"
    )

    # 파일 존재 확인 (Local only)
    if not is_s3_file and not os.path.exists(doc.file_path):
        raise HTTPException(status_code=404, detail="File not found on server")

    # 미디어 타입 추론
    media_type, _ = mimetypes.guess_type(doc.file_path)
    if not media_type:
        media_type = "application/octet-stream"

    # Excel/CSV/Word 파일은 HTML로 변환하여 미리보기 제공
    ext = os.path.splitext(doc.filename)[1].lower()
    if ext in [".xlsx", ".xls", ".csv", ".docx"]:
        temp_file_path = None
        try:
            target_path = doc.file_path

            # S3 파일인 경우 임시 다운로드
            if is_s3_file:
                if doc.file_path.startswith("s3://"):
                    # s3:// 프로토콜은 presigned url 변환이 필요하나, 현재는 http url을 가정
                    pass
                else:
                    response = requests.get(doc.file_path, stream=True)
                    response.raise_for_status()

                    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
                        for chunk in response.iter_content(chunk_size=8192):
                            tmp.write(chunk)
                        temp_file_path = tmp.name
                        target_path = temp_file_path

            body_content = ""

            if ext == ".docx":
                # WORD 처리
                doc_word = DocxDocument(target_path)
                paragraphs = [
                    f"<p>{p.text}</p>" for p in doc_word.paragraphs if p.text.strip()
                ]

                # 표 내용도 간단히 추가
                for table in doc_word.tables:
                    rows_html = []
                    for row in table.rows:
                        cells = [f"<td>{cell.text}</td>" for cell in row.cells]
                        rows_html.append(f"<tr>{''.join(cells)}</tr>")
                    if rows_html:
                        paragraphs.append(
                            f"<table class='docx-table'>{''.join(rows_html)}</table>"
                        )

                body_content = "\n".join(paragraphs)

            else:
                # EXCEL/CSV 처리
                if ext == ".csv":
                    df = pd.read_csv(target_path, nrows=100)
                else:
                    df = pd.read_excel(target_path, nrows=100)

                body_content = f"""
                <div class="info-banner">
                    <span>⚠️</span>
                    성능을 위해 상위 100행만 미리보기로 제공됩니다.
                </div>
                {df.to_html(index=False, border=0)}
                """

            # 공통 HTML 스타일링
            html_content = f"""
            <html>
            <head>
                <style>
                    body {{ font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Helvetica, Arial, sans-serif; margin: 0; padding: 20px; background-color: #ffffff; line-height: 1.6; }}
                    /* Table Styles */
                    table {{ border-collapse: collapse; width: 100%; font-size: 14px; border: 1px solid #e5e7eb; margin-bottom: 20px; }}
                    th {{ background-color: #f9fafb; color: #374151; font-weight: 600; text-align: left; padding: 12px 16px; border-bottom: 1px solid #e5e7eb; }}
                    td {{ padding: 12px 16px; border-bottom: 1px solid #e5e7eb; color: #4b5563; }}
                    tr:last-child td {{ border-bottom: none; }}
                    tr:hover td {{ background-color: #f9fafb; }}
                    
                    /* Docx Specific */
                    p {{ margin-bottom: 0.8em; color: #1f2937; }}
                    .docx-table td {{ border: 1px solid #e5e7eb; }}

                    .info-banner {{
                        margin-bottom: 16px; padding: 10px 14px; background: #fffbeb; border: 1px solid #fcd34d;
                        color: #92400e; border-radius: 6px; font-size: 13px; font-weight: 500; display: flex; align-items: center; gap: 6px;
                    }}
                </style>
            </head>
            <body>
                {body_content}
            </body>
            </html>
            """
            return HTMLResponse(content=html_content)
        except Exception as e:
            logger.error(f"Excel conversion failed: {e}")
            # 변환 실패 시 다운로드로 fallback
        finally:
            if temp_file_path and os.path.exists(temp_file_path):
                try:
                    os.remove(temp_file_path)
                except Exception:
                    pass

    # 브라우저가 s3에서 파일을 직접 받아온다.
    if is_s3_file:
        try:
            # 1. 서버가 S3에서 파일 스트림을 가져옴
            external_res = requests.get(doc.file_path, stream=True)
            external_res.raise_for_status()

            # 2. 클라이언트에게 스트리밍 전송 함수 정의
            def iterfile():
                yield from external_res.iter_content(chunk_size=8192)

            # 3. StreamingResponse 반환
            return StreamingResponse(
                iterfile(),
                media_type=media_type,
                headers={
                    "Content-Disposition": f"inline; filename={requests.utils.quote(doc.filename)}"
                },
            )
        except Exception as e:
            logger.error(f"Failed to proxy S3 file: {e}")
            # 실패 시 Fallback (혹은 에러처리)
            return RedirectResponse(url=doc.file_path)

    return FileResponse(
        doc.file_path,
        filename=doc.filename,
        media_type=media_type,
        content_disposition_type="inline",
    )


@router.post(
    "/{kb_id}/documents/{document_id}/process", status_code=status.HTTP_202_ACCEPTED
)
async def process_document(
    kb_id: UUID,
    document_id: UUID,
    request: DocumentPreviewRequest,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    문서 설정(청킹 등)을 저장하고 백그라운드 처리를 시작합니다.
    """

    # 1. 문서 조회 (권한 확인)
    doc = (
        db.query(Document)
        .join(KnowledgeBase)
        .filter(
            Document.id == document_id,
            KnowledgeBase.id == kb_id,
            KnowledgeBase.user_id == current_user.id,
        )
        .first()
    )

    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # 2. 설정 업데이트
    doc.chunk_size = request.chunk_size
    doc.chunk_overlap = request.chunk_overlap

    # 메타데이터에 추가 설정 저장
    new_meta = dict(doc.meta_info or {})
    new_meta.update(
        {
            "segment_identifier": request.segment_identifier,
            "remove_urls_emails": request.remove_urls_emails,
            "remove_whitespace": request.remove_whitespace,
            "db_config": request.db_config,
            # 필터링 설정 저장
            "selection_mode": request.selection_mode,
            "chunk_range": request.chunk_range,
            "keyword_filter": request.keyword_filter,
        }
    )
    doc.meta_info = new_meta

    # DB 소스인 경우 FK 관계 검증 (백그라운드 실행 전)
    if doc.source_type == "DB" and request.db_config:
        selections = request.db_config.get("selections", [])
        join_config = request.db_config.get("join_config", {})
        
        # 2개 테이블 선택 시 FK 관계 필수
        if len(selections) == 2 and not join_config.get("enabled", False):
            raise HTTPException(
                status_code=400,
                detail="선택한 테이블 간 FK 관계가 없습니다."
            )

    # 상태 업데이트 (처리 시작 전)
    doc.status = (
        "indexing"  # IngestionService가 실행되기 전부터 UI에서 처리중으로 표시하기 위함
    )
    db.commit()

    # 3. 백그라운드 작업 시작
    ingestion_service = IngestionService(
        db,
        user_id=current_user.id,
        chunk_size=request.chunk_size,
        chunk_overlap=request.chunk_overlap,
        ai_model=doc.knowledge_base.embedding_model,
    )

    background_tasks.add_task(
        ingestion_service.process_document,
        document_id,
    )

    return {"status": "processing", "message": "Document processing started"}


@router.post(
    "/{kb_id}/documents/{document_id}/preview", response_model=DocumentPreviewResponse
)
def preview_document_chunking(
    kb_id: UUID,
    document_id: UUID,
    request: DocumentPreviewRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    문서 청킹 설정을 미리보기 합니다. DB를 업데이트하지 않고 결과만 반환합니다.
    """
    # 1. 문서 존재 및 권한 확인
    doc = (
        db.query(Document)
        .join(KnowledgeBase)
        .filter(Document.id == document_id, KnowledgeBase.user_id == current_user.id)
        .first()
    )

    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    if doc.knowledge_base_id != kb_id:
        raise HTTPException(
            status_code=400, detail="Document does not belong to this Knowledge Base"
        )

    # 2. 서비스 호출
    service = IngestionService(db, user_id=current_user.id)
    try:
        segments = service.preview_chunking(
            file_path=doc.file_path,
            chunk_size=request.chunk_size,
            chunk_overlap=request.chunk_overlap,
            segment_identifier=request.segment_identifier,
            remove_urls_emails=request.remove_urls_emails,
            remove_whitespace=request.remove_whitespace,
            strategy=request.strategy,
            source_type=request.source_type,
            meta_info=doc.meta_info,
            db_config=request.db_config,
            # 필터링 파라미터 전달
            selection_mode=request.selection_mode,
            chunk_range=request.chunk_range,
            keyword_filter=request.keyword_filter,
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.exception("Preview failed")
        raise HTTPException(status_code=500, detail=str(e))

    # 3. 응답 반환
    return DocumentPreviewResponse(
        segments=segments,
        total_count=len(segments),
        preview_text_sample="",  # 필요시 원본 텍스트 일부 반환 가능
    )


@router.post(
    "/{kb_id}/documents/{document_id}/sync", status_code=status.HTTP_202_ACCEPTED
)
async def sync_document(
    kb_id: UUID,
    document_id: UUID,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    문서를 동기화합니다. (API 소스 등 재위)
    기존 설정을 유지하면서 처리를 다시 시작합니다.
    """
    # 1. 문서 조회
    doc = (
        db.query(Document)
        .join(KnowledgeBase)
        .filter(
            Document.id == document_id,
            KnowledgeBase.id == kb_id,
            KnowledgeBase.user_id == current_user.id,
        )
        .first()
    )

    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # 상태 업데이트
    doc.status = "indexing"
    db.commit()

    # 2. 백그라운드 작업 시작
    ingestion_service = IngestionService(
        db,
        user_id=current_user.id,
        chunk_size=doc.chunk_size,
        chunk_overlap=doc.chunk_overlap,
        ai_model=doc.knowledge_base.embedding_model,
    )

    background_tasks.add_task(
        ingestion_service.process_document,
        document_id,
    )

    return {"status": "processing", "message": "Document sync started"}
