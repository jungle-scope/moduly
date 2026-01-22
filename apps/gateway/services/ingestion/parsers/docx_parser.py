from typing import Any, Dict, List

from docx import Document as DocxDocument

from apps.gateway.services.ingestion.parsers.base import BaseParser


class DocxParser(BaseParser):
    """
    [DocxParser]
    Word(.docx) 파일을 처리하여 텍스트를 추출합니다.
    본문 텍스트와 표(Table) 내용을 간단한 마크다운/텍스트 형태로 변환합니다.
    """

    def parse(self, source_path: str, **kwargs) -> List[Dict[str, Any]]:
        try:
            doc = DocxDocument(source_path)
            full_text = []

            # 1. 문단(Paragraph) 추출
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)

            # 2. 표(Table) 추출
            # 간단하게 '|' 로 구분된 텍스트로 변환
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    # 빈 행은 제외할 수도 있으나, 표 구조 유지를 위해 포함
                    full_text.append(" | ".join(row_text))

            combined_text = "\n".join(full_text)

            # Word 문서는 페이지 구분이 명확하지 않으므로 전체를 1페이지로 취급
            return [{"text": combined_text, "page": 1}]

        except Exception:
            # 에러 발생 시 빈 리스트 반환 (Processor에서 처리)
            return []
