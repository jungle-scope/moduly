import hashlib
import json
import os  # 폴더 만들기용
import re
import shutil  # 파일 복사용
from enum import Enum
from typing import Optional
from uuid import UUID

import fitz  # PyMuPDF
import pandas as pd
import pymupdf4llm
import tiktoken
from docx import Document as DocxDocument
from fastapi import UploadFile
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from sqlalchemy.orm import Session

from db.models.knowledge import Document, DocumentChunk, SourceType
from db.models.llm import LLMCredential, LLMProvider
from services.data_sources import ApiDataSource, BaseDataSource, FileDataSource


class ParsingStrategy(str, Enum):
    TEXT = "text"
    MIXED = "mixed"
    IMAGE = "image"


class IngestionService:
    def __init__(
        self,
        db: Session,
        user_id: Optional[UUID] = None,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
        ai_model: str = "text-embedding-3-small",
    ):
        self.db = db
        self.user_id = user_id
        self.ai_model = ai_model

        # 청킹 전략 설정
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            # 문단 바뀔 때, 줄 바꿀 때, 마침표, 띄어쓰기일 때 자른다
            separators=["\n\n", "\n", ".", " ", ""],
            keep_separator=True,
        )

    def _get_data_source(self, source_type: str) -> BaseDataSource:
        if source_type == SourceType.FILE:
            return FileDataSource()
        elif source_type == SourceType.API:
            return ApiDataSource()
        # Default fallback for legacy data or if source_type is string "FILE"
        if str(source_type) == "FILE":
            return FileDataSource()
        raise ValueError(f"Unknown source type: {source_type}")

    def save_temp_file(self, file: UploadFile) -> str:
        """
        설명: 메모리에 있는 업로드 파일을 디스크(uploads 폴더)에 저장합니다.
        동일한 파일명이 업로드되어도 물리적 충돌을 방지하기 위해 UUID를 붙여서 저장합니다.
        """
        import uuid

        upload_dir = "uploads"
        os.makedirs(upload_dir, exist_ok=True)

        # 고유한 파일명 생성 (예: a1b2c3d4..._보고서.pdf)
        unique_filename = f"{uuid.uuid4()}_{file.filename}"

        # 저장될 파일의 전체 주소 (예: "uploads/a1b2c3d4..._보고서.pdf")
        file_path = os.path.join(upload_dir, unique_filename)

        with open(file_path, "wb") as buffer:
            # 메모리에 있는 파일(file.file)을 하드디스크(buffer)로 복사
            shutil.copyfileobj(file.file, buffer)

        return file_path

    def create_pending_document(
        self,
        knowledge_base_id: UUID,
        filename: str,
        file_path: str | None,
        chunk_size: int,
        chunk_overlap: int,
        source_type: SourceType = SourceType.FILE,
        meta_info: dict = None,
    ) -> UUID:
        """
        파일 업로드 시점에 'Pending' 상태의 Document 레코드를 먼저 생성합니다.
        KnowledgeBase와의 연결(FK)을 위해 knowledge_base_id가 필수입니다.
        설정된 chunk_size와 chunk_overlap을 저장하여 나중에 참조할 수 있게 합니다.
        """
        new_doc = Document(
            knowledge_base_id=knowledge_base_id,
            filename=filename,
            file_path=file_path,
            status="pending",
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            source_type=source_type,
            meta_info=meta_info or {},
        )
        self.db.add(new_doc)
        self.db.commit()
        self.db.refresh(new_doc)
        return new_doc.id

    async def process_document_background(
        self, document_id: UUID, knowledge_base_id: UUID, file_path: str
    ):
        """
        BackgroundTasks의 메인 진입점.
        파싱 -> 청킹 -> 임베딩 -> 저장
        """
        try:
            self._update_status(document_id, "indexing")
            self._update_progress(document_id, 5, "문서 처리를 시작합니다...")
            print("[DEBUG] 1번")
            # 1단계: 파싱 (document_id 전달)
            self._update_progress(document_id, 10, "문서 내용을 분석하고 있습니다...")

            # 1.3 DataSource를 통한 텍스트 추출
            # DB에서 Document 객체 조회
            doc = self.db.query(Document).get(document_id)
            if not doc:
                raise ValueError(f"Document {document_id} not found")

            data_source = self._get_data_source(doc.source_type)
            print("[DEBUG] 2번", data_source)

            # 소스 설정 구성
            source_config = {}
            if doc.source_type == SourceType.FILE or str(doc.source_type) == "FILE":
                source_config = {
                    "file_path": file_path,
                    "document_id": str(document_id),
                }
            elif doc.source_type == SourceType.API:
                # meta_info에서 API 설정 가져오기
                api_config = doc.meta_info.get("api_config", {})

                # 헤더 복호화 로직
                import json

                from core.security import security_service

                headers = api_config.get("headers")
                if headers and isinstance(headers, str):
                    try:
                        decrypted_json = security_service.decrypt(headers)
                        api_config["headers"] = json.loads(decrypted_json)
                    except Exception as e:
                        print(f"Failed to decrypt headers: {e}")
                        # 복호화 실패 시 빈 딕셔너리 사용하거나 에러 처리
                        api_config["headers"] = {}

                source_config = api_config

            text_blocks = data_source.fetch_text(source_config)

            # 파싱 결과가 비어있다면 (비용 승인 대기 등) 중단
            if not text_blocks:
                doc = self.db.query(Document).get(document_id)
                if doc and doc.status == "waiting_for_approval":
                    print(f"⏸️ Document {document_id} paused for approval.")
                    self._update_progress(
                        document_id, 0, "추가 비용 승인이 필요하여 대기 중입니다."
                    )
                    return
                # 진짜 내용이 없는 경우일 수도 있음 (이 경우 completed 처리됨)
                print(f"⚠️ No text extracted from document {document_id}")
                self._update_status(document_id, "completed")
                return

            # 1.5 Content Hash Check (변경 감지)
            # 모든 텍스트를 합쳐서 해시 생성
            full_text = "".join([b["text"] for b in text_blocks])
            new_hash = hashlib.sha256(full_text.encode("utf-8")).hexdigest()

            if doc.content_hash == new_hash:
                print(f"⏭️ Content unchanged for {document_id}. Skipping processing.")
                self._update_progress(
                    document_id, 100, "변경된 내용이 없어 처리를 건너뜁니다."
                )
                self._update_status(document_id, "completed")
                return

            # 해시 업데이트
            doc = self.db.query(Document).get(document_id)
            if doc:
                doc.content_hash = new_hash
                self.db.commit()

            self._update_progress(document_id, 40, "문서 내용 분석이 완료되었습니다.")

            # 2~4단계: 청킹, 임베딩, 저장 및 완료 처리
            self._finalize_ingestion(document_id, knowledge_base_id, text_blocks)
        except Exception as e:
            print(f"Ingestion failed: {e}")
            self._update_status(document_id, "failed", error_message=str(e))
            self._update_progress(
                document_id, 0, f"처리 중 오류가 발생했습니다: {str(e)}"
            )

    async def resume_processing(self, document_id: UUID, strategy: str = "llamaparse"):
        """
        승인된 문서에 대해 파싱을 재개합니다.
        strategy: 'llamaparse' (유료, 고품질) or 'general' (무료, PyMuPDF)
        """
        doc = self.db.query(Document).get(document_id)
        if not doc:
            print(f"❌ Document {document_id} not found for resumption.")
            return

        try:
            print(f"▶️ Resuming ingestion for {document_id} with strategy: {strategy}")
            self._update_status(document_id, "indexing")

            text_blocks = []

            # 1단계: 전략에 따른 파싱 (DataSource 사용)
            data_source = self._get_data_source(doc.source_type)

            source_config = {}
            if doc.source_type == SourceType.FILE:
                source_config = {
                    "file_path": doc.file_path,
                    "document_id": str(document_id),
                    "strategy": strategy,  # "general" or "llamaparse" passed from arg
                }
            elif doc.source_type == SourceType.API:
                api_config = doc.meta_info.get("api_config", {})
                source_config = api_config

            text_blocks = data_source.fetch_text(source_config)

            # 2~4단계: 청킹, 임베딩, 저장 및 완료 처리
            self._finalize_ingestion(document_id, doc.knowledge_base_id, text_blocks)

        except Exception as e:
            print(f"❌ Resumption failed: {e}")
            self._update_status(document_id, "failed", error_message=str(e))

    def _analyze_pdf_type(self, file_path: str) -> str:
        """
        PDF 파일의 성격을 파악하여 적절한 파싱 전략을 반환합니다.

        Sampling Strategy:
        - 앞 3페이지 + 중간 1페이지 + 뒤 2페이지 (총 최대 6페이지)

        Returns:
            - 'special': 전체가 이미지거나 텍스트가 거의 없는 경우 (LlamaParse 등 필요) -> OCR 필요
            - 'fast': 텍스트 위주의 일반적인 문서 (PyMuPDF4LLM 사용)
            - 'precise': 텍스트와 이미지가 섞여있어 정밀한 레이아웃 분석이 필요한 경우
        """
        doc = fitz.open(file_path)
        total_pages = len(doc)

        # 1. 너무 큰 파일 예외 처리 (예: 500페이지 이상은 일단 경고)
        if total_pages > 500:
            print(f"[Warn] Large file detected: {total_pages} pages.")

        # 2. 샘플링 페이지 인덱스 선정
        sample_indices = set()

        # 앞 3페이지
        for i in range(min(3, total_pages)):
            sample_indices.add(i)

        # 중간 1페이지
        if total_pages > 3:
            sample_indices.add(total_pages // 2)

        # 뒤 2페이지
        if total_pages > 1:
            sample_indices.add(total_pages - 1)
        if total_pages > 2:
            sample_indices.add(total_pages - 2)

        sorted_indices = sorted(list(sample_indices))

        # 3. 샘플링 분석
        image_count = 0
        text_length = 0
        page_count = 0

        for idx in sorted_indices:
            if idx >= total_pages:
                continue

            page = doc[idx]
            page_count += 1

            # 텍스트 추출
            text = page.get_text()
            text_length += len(text.strip())

            # 이미지 객체 카운트
            images = page.get_images(full=True)
            image_count += len(images)

        doc.close()

        # 4. 분석 결과에 따른 전략 결정

        # 평균 텍스트 길이 (페이지당)
        avg_text_per_page = text_length / page_count if page_count > 0 else 0

        # 평균 이미지 수 (페이지당)
        avg_images_per_page = image_count / page_count if page_count > 0 else 0

        print(
            f"[PDF Analysis] Avg Text: {avg_text_per_page:.1f}, Avg Images: {avg_images_per_page:.1f}"
        )

        # Case A: 텍스트가 거의 없음 (OCR 필요)
        if avg_text_per_page < 50:
            return ParsingStrategy.IMAGE

        # Case B: 이미지가 많고 텍스트도 어느정도 있음 (복잡한 레이아웃 가능성)
        elif avg_images_per_page > 2:
            return ParsingStrategy.MIXED

        # Case C: 텍스트 위주
        else:
            return ParsingStrategy.TEXT

    def _parse_with_pymupdf(self, file_path: str) -> list[dict]:
        """기존 PyMuPDF4LLM 기반 파싱 로직"""
        md_text_chunks = pymupdf4llm.to_markdown(file_path, page_chunks=True)

        results = []
        for chunk in md_text_chunks:
            results.append(
                {
                    "text": chunk["text"],
                    "page": chunk["metadata"]["page"] + 1,
                }
            )
        return results

    def _get_cache_path(self, file_path: str) -> str:
        """캐시 파일 경로 생성 (원본파일_parsed.json)"""
        return f"{file_path}_parsed.json"

    def _save_cache(self, file_path: str, data: list[dict]):
        """파싱 결과를 JSON으로 저장"""
        try:
            cache_path = self._get_cache_path(file_path)
            with open(cache_path, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            print(f"💾 [Cache Saved] {cache_path}")
        except Exception as e:
            print(f"⚠️ Failed to save cache: {e}")

    def _load_cache(self, file_path: str) -> list[dict]:
        """캐시된 파싱 결과 로드"""
        cache_path = self._get_cache_path(file_path)
        if os.path.exists(cache_path):
            try:
                with open(cache_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                print(f"♻️ [Cache Hit] Loaded parsing result from {cache_path}")
                return data
            except Exception as e:
                print(f"⚠️ Cache load failed: {e}")
        return None

    def _get_llamaparse_api_key(self) -> Optional[str]:
        """
        LlamaParse API 키를 가져옵니다.
        1) 환경변수 우선
        2) 현재 사용자(user_id)의 LlamaParse 자격증명
        3) 사용자 미지정 시 시스템 내 첫 번째 유효 자격증명
        """
        env_key = os.getenv("LLAMA_CLOUD_API_KEY")
        if env_key:
            return env_key

        provider = (
            self.db.query(LLMProvider).filter(LLMProvider.name == "llamaparse").first()
        )
        if not provider:
            return None

        query = self.db.query(LLMCredential).filter(
            LLMCredential.provider_id == provider.id,
            LLMCredential.is_valid == True,
        )
        if self.user_id:
            query = query.filter(LLMCredential.user_id == self.user_id)

        cred = query.order_by(LLMCredential.created_at.desc()).first()

        # 사용자별 키가 없으면 시스템 내 아무 유효 키나 사용 (환경변수 접근과 동일 레벨)
        if not cred:
            cred = (
                self.db.query(LLMCredential)
                .filter(
                    LLMCredential.provider_id == provider.id,
                    LLMCredential.is_valid == True,
                )
                .order_by(LLMCredential.created_at.desc())
                .first()
            )

        if not cred:
            return None

        try:
            cfg = json.loads(cred.encrypted_config)
            return cfg.get("apiKey")
        except Exception as e:
            print(f"[Warning] Failed to parse LlamaParse credential: {e}")
            return None

    def _parse_with_llamaparse(self, file_path: str) -> list[dict]:
        """LlamaParse API 연동 (캐싱 적용)"""

        # 1. 캐시 확인
        cached_data = self._load_cache(file_path)
        if cached_data:
            return cached_data

        # 비용 예측 로그 출력
        est = self._estimate_llamaparse_cost(file_path)
        print(
            f"💰 [비용 예측] 페이지 수: {est['pages']}, 크레딧: {est['credits']}, 비용: ${est['cost_usd']:.4f}"
        )

        try:
            from llama_parse import LlamaParse
        except ImportError:
            print(
                "❌ LlamaParse 라이브러리를 찾을 수 없습니다. 'pip install llama-parse'를 실행해주세요."
            )
            return []

        api_key = self._get_llamaparse_api_key()
        if not api_key:
            print(
                "❌ LlamaParse API Key가 설정되지 않았습니다. "
                "환경변수 LLAMA_CLOUD_API_KEY를 설정하거나 설정 > LLM Provider에서 키를 등록해주세요."
            )
            return []

        print("🚀 LlamaParse 클라우드 처리 시작...")

        try:
            # 파서 초기화
            # result_type="markdown"이 기본값이지만 명시적으로 설정
            # language="ko"를 설정하여 한국어 인식률 향상
            parser = LlamaParse(
                api_key=api_key,
                result_type="markdown",
                verbose=True,
                language="ko",
                fast_mode=True,
            )

            # JSON 결과를 받아야 페이지별 텍스트와 메타데이터를 확실하게 구분할 수 있음
            # get_json_result는 파일당 하나의 결과 객체를 리스트로 반환함
            json_results = parser.get_json_result(file_path)

            # [Debug] 구조 확인
            # print(f"🔍 [LlamaParse Raw Result]: {json_results}")

            parsed_results = []
            full_text_for_debug = ""

            if json_results and isinstance(json_results, list):
                first_result = json_results[0]
                # 'pages' 키에 각 페이지별 파싱 결과가 담겨있음
                pages = first_result.get("pages", [])

                for p in pages:
                    # 'md' 키가 없을 경우를 대비해 키 확인
                    # 'md' 키가 없으면 'text' 키를 사용 (fast_mode 등에서 발생)
                    md_text = p.get("md") or p.get("text") or ""
                    parsed_results.append(
                        {
                            "text": md_text,  # 마크다운 변환 텍스트
                            "page": p["page"],  # 페이지 번호
                        }
                    )
                    full_text_for_debug += f"\n--- Page {p['page']} ---\n{md_text}\n"

            # [Debug] 파싱된 결과를 파일로 저장
            # try:
            #     base_dir = os.path.dirname(file_path)
            #     file_name = os.path.basename(file_path)
            #     debug_file_name = f"{os.path.splitext(file_name)[0]}_parsed.md"
            #     debug_file_path = os.path.join(base_dir, debug_file_name)

            #     with open(debug_file_path, "w", encoding="utf-8") as f:
            #         f.write(full_text_for_debug)
            #     print(f"💾 [Debug] Parsed content saved to: {debug_file_path}")
            # except Exception as e:
            #     print(f"⚠️ Failed to save debug file: {e}")

            print(f"LlamaParse 완료: 총 {len(parsed_results)} 페이지 변환됨.")

            # 2. 결과 캐싱
            self._save_cache(file_path, parsed_results)

            return parsed_results

        except Exception as e:
            print(f"LlamaParse 처리 실패: {e}")
            return []

    def _estimate_llamaparse_cost(self, file_path: str) -> dict:
        """
        LlamaParse 예측 비용 계산
        기준: Standard Mode (3 credits/page), $1 = 1000 credits
        """
        try:
            # PDF가 아닌 경우 fitz.open()이 실패할 수 있으므로 예외 처리
            ext = os.path.splitext(file_path)[1].lower()
            if ext not in [".pdf", ".xps", ".epub", ".mobi", ".fb2", ".cbz", ".svg"]:
                # 이미지 파일(png, jpg 등)은 fitz로 열 수 있지만, 엑셀/워드는 불가
                # 일단 0으로 리턴하여 fallback 유도
                if ext not in [".png", ".jpg", ".jpeg", ".tiff", ".bmp"]:
                    return {"pages": 0, "credits": 0, "cost_usd": 0.0}

            doc = fitz.open(file_path)
            total_pages = len(doc)
            doc.close()

            # Standard Mode 기준 (페이지당 3 크레딧)
            credits_per_page = 3
            total_credits = total_pages * credits_per_page
            cost_usd = total_credits / 1000.0

            return {
                "pages": total_pages,
                "credits": total_credits,
                "cost_usd": cost_usd,
            }
        except Exception as e:
            print(f"[Warning] Cost estimation failed: {e}")
            return {"pages": 0, "credits": 0, "cost_usd": 0.0}

    def _parse_excel_csv(self, file_path: str) -> list[dict]:
        """Excel/CSV 파일 파싱 (모든 시트 처리)"""
        text_content = ""
        ext = os.path.splitext(file_path)[1].lower()

        try:
            if ext == ".csv":
                df = pd.read_csv(file_path)
                text_content += f"# CSV Content\n\n{df.to_markdown(index=False)}\n"
            else:
                # sheet_name=None -> 모든 시트를 dict로 반환
                xls = pd.read_excel(file_path, sheet_name=None)
                for sheet_name, df in xls.items():
                    text_content += f"\n# Sheet: {sheet_name}\n\n"
                    text_content += df.to_markdown(index=False) + "\n"

            # 엑셀은 페이지 개념이 모호하므로 전체를 1페이지로 취급하거나 적절히 분할
            return [{"text": text_content, "page": 1}]
        except Exception as e:
            print(f"Excel/CSV parsing failed: {e}")
            return []

    def _parse_docx(self, file_path: str) -> list[dict]:
        """Word(.docx) 파일 파싱"""
        try:
            doc = DocxDocument(file_path)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)

            # 간단한 표 처리
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    full_text.append(" | ".join(row_text))

            return [{"text": "\n".join(full_text), "page": 1}]
        except Exception as e:
            print(f"Docx parsing failed: {e}")
            return []

    def _parse_txt(self, file_path: str) -> list[dict]:
        """Text/Markdown 파일 파싱"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            return [{"text": content, "page": 1}]
        except Exception as e:
            print(f"Text parsing failed: {e}")
            return []

    def _is_mixed_quality_poor(self, results: list[dict]) -> bool:
        """MIXED 모드 품질 검사: 레이아웃이 심각하게 깨졌는지 확인"""
        total_text = "".join([r["text"] for r in results])

        # 휴리스틱 1: 알 수 없는 특수문자나 공백 패턴이 너무 많은 경우
        if len(total_text) > 0:
            broken_char_count = total_text.count("\ufffd")
            if (broken_char_count / len(total_text)) > 0.05:  # 5% 이상 깨짐
                print("Reason: High broken character rate in MIXED mode.")
                return True

        # 휴리스틱 2: 마크다운 구조가 거의 없음 (헤더 #이 너무 적음)
        # 일반적인 문서라면 페이지당 적어도 1~2개의 헤더는 있어야 함
        page_count = len(results)
        header_count = total_text.count("\n#")
        if (
            page_count > 0 and (header_count / page_count) < 0.2
        ):  # 5페이지당 헤더 1개 미만
            print("Reason: Too few markdown headers found.")
            return True

        return False

    def _is_text_quality_poor(self, file_path: str, results: list[dict]) -> bool:
        """TEXT 모드 품질 검사"""
        total_text = "".join([r["text"] for r in results])

        # 1. 글자 수가 너무 적음 (50자 미만)
        if len(total_text.strip()) < 50:
            print("Reason: Too few characters extracted.")
            return True

        # 2. 깨진 문자(replacement character ) 비율 확인
        broken_char_count = total_text.count("\ufffd")  # or other garbage chars
        if len(total_text) > 0 and (
            broken_char_count / len(total_text) > 0.05
        ):  # 5% 이상
            print("Reason: Too many broken characters.")
            return True

        # 3. (고급) PyMuPDF로 표(Table)는 감지되는데, 추출된 텍스트에는 마크다운 표 문법(|---|)이 없는 경우
        try:
            doc = fitz.open(file_path)
            has_table_but_no_markdown = False

            # 성능을 위해 앞부분 5페이지만 검사
            for i in range(min(5, len(doc))):
                page = doc[i]
                tables = page.find_tables()
                if tables and len(tables.tables) > 0:
                    # 해당 페이지의 추출된 텍스트 찾기
                    page_text = results[i]["text"] if i < len(results) else ""
                    # 표는 있는데 마크다운 표 구문('|')이 전혀 없다면 파싱 실패로 간주
                    if "|" not in page_text:
                        print(
                            f"Reason: Table detected on page {i + 1} but no markdown table found."
                        )
                        has_table_but_no_markdown = True
                        break

            doc.close()
            if has_table_but_no_markdown:
                return True

        except Exception as e:
            print(f"[Warning] Table check failed: {e}")
            # 에러 나면 안전하게 False 반환 (Flow 중단 안 함)
            return False

        return False

    def _request_llamaparse_approval(
        self, file_path: str, document_id: UUID
    ) -> list[dict]:
        """
        LlamaParse 호출 전 비용 계산 후 '승인 대기' 상태로 변경하고 중단함.
        """
        if not document_id:
            # document_id가 없으면(디버그/테스트 모드) 그냥 진행
            print("No document_id provided. Skipping approval and running LlamaParse.")
            return self._parse_with_llamaparse(file_path)

        # 1. 비용 계산
        est = self._estimate_llamaparse_cost(file_path)

        # 2. DB 업데이트 (상태: waiting_for_approval)
        doc = self.db.query(Document).get(document_id)
        if doc:
            doc.status = "waiting_for_approval"
            # 기존 메타데이터에 비용 정보 병합
            new_meta = dict(doc.meta_info or {})
            new_meta.update({"cost_estimate": est, "strategy": "llamaparse_fallback"})
            doc.meta_info = new_meta
            self.db.commit()

        print(
            f"⏸️ [Approval Required] Document {document_id} paused for LlamaParse cost approval."
        )

        # 3. 빈 리스트 반환하여 파이프라인 중단
        return []

    def _parse_pdf(self, file_path: str, document_id: UUID = None) -> list[dict]:
        """
        PDF 파싱 메인 진입점.
        적절한 파서(PyMuPDF / LlamaParse)를 선택하고,
        품질 저하 시 Fallback 로직을 수행 (비용 승인 프로세스 포함)
        """
        # 1. 파일 성격 파악
        parsing_strategy = self._analyze_pdf_type(file_path)
        print(f"[{file_path}] Parsing Strategy: {parsing_strategy.value}")

        # Case 1: 이미지 위주 (OCR 필수) -> 승인 요청
        if parsing_strategy == ParsingStrategy.IMAGE:
            print("Strategy is IMAGE. Requesting approval for LlamaParse.")
            return self._request_llamaparse_approval(file_path, document_id)

        # Case 2: 혼합형 (텍스트 + 이미지)
        elif parsing_strategy == ParsingStrategy.MIXED:
            # 1차 시도: PyMuPDF (빠름)
            results = self._parse_with_pymupdf(file_path)

            # 품질 검사: 결과물이 '난잡'한지 확인
            if self._is_mixed_quality_poor(results):
                print(
                    "Mixed parsing quality is poor. Requesting approval for LlamaParse."
                )
                return self._request_llamaparse_approval(file_path, document_id)

            return results

        # Case 3: 텍스트 위주
        else:  # ParsingStrategy.TEXT
            # 1차 시도: PyMuPDF
            results = self._parse_with_pymupdf(file_path)

            # 품질 검사: 텍스트 누락, 깨짐, 표 구조 이상 확인
            if self._is_text_quality_poor(file_path, results):
                print(
                    "Text parsing quality is poor. Requesting approval for LlamaParse."
                )
                return self._request_llamaparse_approval(file_path, document_id)

            return results

    def _create_chunks(self, text_blocks: list[dict]) -> list[dict]:
        """
        파싱된 텍스트를 더 작은 조각(Chunk)으로 나눕니다.
        """
        final_chunks = []

        for block in text_blocks:
            splits = self.text_splitter.split_text(block["text"])
            for split in splits:
                final_chunks.append(
                    {"content": split, "metadata": {"page": block["page"]}}
                )
        return final_chunks

    def _save_chunks_to_pgvector(
        self, document_id: UUID, knowledge_base_id: UUID, chunks: list[dict]
    ):
        """
        텍스트 조각들을 OpenAI에 보내서 '의미 벡터'로 바꾼 뒤, DocumentChunk 테이블에 저장합니다.
        기존 청크가 있다면 삭제하고 새로 저장합니다 (Clean & Insert).
        """
        print(f"🔍 [Debug] _save_chunks_to_pgvector 시작: doc_id={document_id}")
        # 0. 기존 청크 삭제 (Clean Step)
        try:
            del_count = (
                self.db.query(DocumentChunk)
                .filter(DocumentChunk.document_id == document_id)
                .delete()
            )
            self.db.commit()
            print(f"🗑️ [Debug] 기존 청크 {del_count}개 삭제 완료")
        except Exception as e:
            print(f"❌ [Debug] 기존 청크 삭제 중 에러: {e}")

        # TODO: 토큰 계산을 위한 인코더 설정
        try:
            encoding = tiktoken.encoding_for_model(self.ai_model)
        except KeyError:
            encoding = tiktoken.get_encoding("cl100k_base")  # gpt-4로 가정하고 계산

        # DB에서 API Key 가져오기 (환경변수 의존 제거)
        from db.models.llm import LLMProvider

        api_key = None

        doc = self.db.query(Document).filter(Document.id == document_id).first()
        if not doc:
            print("❌ [Debug] 문서를 찾을 수 없음")
            raise ValueError("문서를 찾을 수 없습니다.")

        user_id = doc.knowledge_base.user_id
        print(f"🔍 [Debug] 문서 소유자 ID: {user_id}")

        user_crd = (
            self.db.query(LLMCredential)
            .join(LLMProvider)
            .filter(
                LLMCredential.user_id == user_id,
                LLMCredential.is_valid,
                LLMProvider.name == "openai",
            )
            .first()
        )

        if user_crd:
            print(f"✅ [Debug] OpenAI 자격 증명 발견 (ID: {user_crd.id})")
            try:
                config = json.loads(user_crd.encrypted_config)
                api_key = config.get("apiKey")
            except Exception as e:
                print(f"[Debug] Credential config 파싱 실패: {e}")

        if not api_key:
            raise ValueError(
                "사용자의 OpenAI API Key를 찾을 수 없습니다. 등록해주세요."
            )
            print("⚠️ [Debug] OpenAI 자격 증명을 찾지 못함")
        print(f"✅ [Debug] API Key 확보 완료 (Key: {api_key[:8]}...)")

        # 임베딩 모델 초기화 (API Key 명시)
        embeddings_model = OpenAIEmbeddings(model=self.ai_model, openai_api_key=api_key)

        # 1. 텍스트 추출
        texts = [chunk["content"] for chunk in chunks]
        print(f"🔍 [Debug] 임베딩 요청 시작 (청크 개수: {len(texts)}개)")

        # 2. 임베딩 생성 (일괄 호출) - 실제 API 사용!
        try:
            embedded_vectors = embeddings_model.embed_documents(texts)
            print("✅ [Debug] 임베딩 생성 완료")
        except Exception as e:
            print(f"❌ [Debug] OpenAI Embedding Error: {e}")
            raise e

        # 3. DB 객체 생성
        try:
            chunk_objects = []
            for i, chunk in enumerate(chunks):
                content = chunk["content"]
                token_count = len(encoding.encode(content))

                db_chunk = DocumentChunk(
                    document_id=document_id,
                    knowledge_base_id=knowledge_base_id,  # 검색 최적화용
                    content=content,
                    embedding=embedded_vectors[i],
                    chunk_index=i,
                    token_count=token_count,
                    metadata_=chunk["metadata"],
                )
                chunk_objects.append(db_chunk)

            print(
                f"📦 [Debug] 저장할 객체 {len(chunk_objects)}개 생성됨. DB에 추가(add) 시도..."
            )
            self.db.add_all(chunk_objects)
            print("💾 [Debug] 커밋(Commit) 시도...")
            self.db.commit()
            print("🎉 [Debug] DB 저장 및 커밋 성공!")

        except Exception as e:
            print(f"❌ [Debug] DB 저장 실패 (Commit Error): {e}")
            self.db.rollback()  # 롤백 시도
            raise e

    def _create_chunks(self, text_blocks: list[dict]) -> list[dict]:
        """
        텍스트 블록 리스트를 받아, 설정된 chunk_size와 chunk_overlap에 따라 청킹합니다.
        각 청크는 원본 텍스트 블록의 메타데이터를 유지하거나 병합할 수 있습니다.
        """
        chunks = []
        for block in text_blocks:
            text = block["text"]
            metadata = block.get("metadata", {})

            # 텍스트가 너무 짧으면 스킵할 수도 있음 (선택사항)
            if not text.strip():
                continue

            splits = self.text_splitter.split_text(text)

            for split in splits:
                chunks.append(
                    {
                        "content": split,
                        "metadata": metadata,  # 페이지 번호 등 원본 메타데이터 보존
                    }
                )

        return chunks

    def _finalize_ingestion(
        self, document_id: UUID, knowledge_base_id: UUID, text_blocks: list[dict]
    ):
        """
        텍스트 블록을 받아 청킹 -> 임베딩 -> 저장 -> 완료 처리를 수행합니다.
        """
        # 2단계: 청킹
        self._update_progress(
            document_id, 50, "AI가 읽기 좋게 문서를 조각내고 있습니다..."
        )
        chunks = self._create_chunks(text_blocks)

        # 3 & 4단계: 임베딩 및 저장
        self._update_progress(
            document_id, 70, "벡터 데이터베이스에 저장할 준비를 하고 있습니다..."
        )
        self._save_chunks_to_pgvector(document_id, knowledge_base_id, chunks)

        # 완료 상태 업데이트
        self._update_progress(document_id, 100, "모든 처리가 완료되었습니다.")
        self._update_status(document_id, "completed")

    def _update_status(self, document_id: UUID, status: str, error_message: str = None):
        doc = self.db.query(Document).get(document_id)
        if doc:
            doc.status = status
            if error_message:
                doc.error_message = error_message
            self.db.commit()

    def _update_progress(self, document_id: UUID, progress: int, message: str):
        """
        문서 처리 진행률(%)과 현재 단계 메시지를 meta_info에 업데이트합니다.
        """
        doc = self.db.query(Document).get(document_id)
        if doc:
            new_meta = dict(doc.meta_info or {})
            new_meta.update(
                {"processing_progress": progress, "processing_current_step": message}
            )
            doc.meta_info = new_meta
            self.db.commit()

    async def analyze_document(self, document_id: UUID) -> dict:
        """
        문서 분석: 페이지 수, 비용 예측 등을 반환
        """
        doc = self.db.query(Document).get(document_id)
        if not doc:
            raise ValueError("Document not found")

        # 1. 비용 예측 (FileDataSource 사용)
        try:
            # 임시로 FILE 타입 가정 (API 등은 0 반환)
            data_source = self._get_data_source(doc.source_type)
            source_config = {}
            if doc.source_type == SourceType.FILE:
                source_config = {"file_path": doc.file_path}

            cost_info = data_source.estimate_cost(source_config)
        except Exception as e:
            print(f"Cost estimation failed: {e}")
            cost_info = {"pages": 0, "credits": 0, "cost_usd": 0.0}

        # 2. 파일 타입 분석 (선택 사항)
        # parsing_strategy = self._analyze_pdf_type(doc.file_path)

        # 3. 캐시 확인 (파일인 경우에만)
        is_cached = False
        cache_path = ""

        if doc.source_type == SourceType.FILE and doc.file_path:
            cache_path = self._get_cache_path(doc.file_path)
            is_cached = os.path.exists(cache_path)

        print(
            f"🔍 [Debug] analyze_document: filename={doc.filename}, is_cached={is_cached}, path={cache_path}"
        )

        return {
            "cost_estimate": cost_info,
            "filename": doc.filename,
            "is_cached": is_cached,
            # "recommended_strategy": parsing_strategy
        }

    def _get_cache_path(self, file_path: str) -> str:
        """
        LlamaParse 결과 캐시 파일 경로를 반환합니다.
        (예: uploads/file.pdf -> uploads/file.pdf.md)
        """
        if not file_path:
            return ""
        return f"{file_path}.md"

    def preview_chunking(
        self,
        file_path: str,
        chunk_size: int,
        chunk_overlap: int,
        segment_identifier: str,
        remove_urls_emails: bool = False,
        remove_whitespace: bool = True,
        strategy: str = "general",  # "general" or "llamaparse",
        source_type: SourceType = SourceType.FILE,
        meta_info: dict = None,
    ) -> list[dict]:
        """
        DB 저장 없이 메모리 상에서 청킹 결과를 미리봅니다.
        strategy에 따라 일반 파싱 또는 정밀 파싱(LlamaParse)을 수행합니다.
        """
        # 1. 텍스트 추출
        try:
            if source_type == SourceType.API:
                # API 반환값 처리, 헤더 복호화
                api_config = meta_info.get("api_config", {})
                headers = api_config.get("headers")
                if headers and isinstance(headers, str):
                    try:
                        from core.security import security_service

                        decrypted_json = security_service.decrypt(headers)
                        api_config["headers"] = json.loads(decrypted_json)
                    except Exception as e:
                        print(f"Failed to decrypt headers: {e}")
                        api_config["headers"] = {}

                data_source = ApiDataSource()
                source_config = api_config
            else:
                data_source = FileDataSource()
                source_config = {
                    "file_path": file_path,
                    "strategy": strategy,  # "general" or "llamaparse"
                }

            text_blocks = data_source.fetch_text(source_config)
            full_text = "\n".join([block["text"] for block in text_blocks])
        except Exception as e:
            print(f"Preview parsing failed: {e}")
            return []

        # 2. 전처리
        if remove_urls_emails:
            # URL 제거
            full_text = re.sub(
                r"http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+",
                "",
                full_text,
            )
            # 이메일 제거
            full_text = re.sub(r"[\w\.-]+@[\w\.-]+", "", full_text)

        if remove_whitespace:
            # 연속된 공백, 탭을 단일 공백으로 치환
            full_text = re.sub(r"[ \t]+", " ", full_text)
            # 연속된 줄바꿈이 3개 이상이면 2개(\n\n)로 축소 (문단 구분 유지)
            full_text = re.sub(r"\n{3,}", "\n\n", full_text)

        # 3. 청킹 설정 오버라이드
        # segment_identifier가 유효하면 separator 목록의 최우선 순위로 추가
        separators = ["\n\n", "\n", ".", " ", ""]
        if segment_identifier and segment_identifier not in separators:
            # 특수 문자(escaped) 처리 필요할 수 있음. 일단 있는 그대로 사용.
            # 사용자가 "\n\n"을 입력하면 문자열 그대로 들어오므로, 실제 이스케이프 시퀀스로 변환해주는 로직이 필요할 수 있음.
            # 프론트에서 실제 줄바꿈을 보내거나, 여기서 변환해야 함.
            # 일단은 단순 문자열 매칭으로 가정하되, \n은 특별 취급
            processed_identifier = segment_identifier.replace("\\n", "\n")
            if processed_identifier not in separators:
                separators.insert(0, processed_identifier)

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=separators,
            keep_separator=True,
        )

        splits = splitter.split_text(full_text)

        # 4. 결과 포맷팅 & 토큰 계산
        try:
            encoding = tiktoken.encoding_for_model(self.ai_model)
        except KeyError:
            encoding = tiktoken.get_encoding("cl100k_base")

        preview_segments = []
        for split in splits:
            token_count = len(encoding.encode(split))
            preview_segments.append(
                {
                    "content": split,
                    "token_count": token_count,
                    "char_count": len(split),
                }
            )

        return preview_segments
